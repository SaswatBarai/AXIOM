"""Phase 13 — Cover Letter Generator: Gemini-powered drafting + PDF/DOCX export."""
from __future__ import annotations

import io
import os
import textwrap
from typing import Literal

import google.generativeai as genai

# ── Gemini setup ──────────────────────────────────────────────────────────────

_API_KEY    = os.getenv("GEMINI_API_KEY", "")
_MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
genai.configure(api_key=_API_KEY)

Tone = Literal["formal", "friendly", "direct"]


def sanitize_input(text: str, max_len: int = 4000) -> str:
    """Strip control characters and enforce length to prevent prompt injection."""
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return cleaned[:max_len]

# ── Few-shot examples (voice anchors) ─────────────────────────────────────────

_FEW_SHOTS = """
Example 1 — Formal tone:
Dear Hiring Manager,
I am writing to express my strong interest in the Senior Backend Engineer role at Stripe. During my four years at Razorpay, I reduced payment processing latency by 38% by re-architecting our distributed cache layer, handling over 2 million transactions daily. My expertise in Go, PostgreSQL, and Kubernetes aligns closely with your infrastructure stack. I would welcome the opportunity to bring this same rigour to Stripe's mission of increasing the GDP of the internet.
Sincerely, [Candidate]

Example 2 — Friendly tone:
Hi [Recruiter Name],
When I saw the Full-Stack Engineer opening at Notion, I immediately knew it was the right fit. I've spent the last three years building collaborative editing tools at Coda, shipping features used by 500k+ teams. I thrive in fast-moving product environments and genuinely love the challenge of making complex workflows feel effortless for users. I'd love to chat about how I can contribute to Notion's next chapter.
Best, [Candidate]

Example 3 — Direct tone:
To the Engineering Team at Linear,
Here's why I'm your next iOS Engineer: I shipped Linear's iOS app at my previous company (Raycast), cut cold-start time by 52%, and reduced crash rate to 0.03%. I know your stack, I've admired your product for years, and I'm ready to make it better. Available to start in 2 weeks.
[Candidate]
"""

# ── System prompt ─────────────────────────────────────────────────────────────

_SYSTEM = """You are an expert career coach and professional writer specializing in tailored cover letters.

Rules:
1. Write in the specified tone (formal / friendly / direct).
2. Use ONLY the concrete metrics and achievements from the candidate's resume — never invent numbers.
3. Open with a compelling hook that references the specific company and role.
4. Highlight 2-3 specific resume achievements that directly match the job description.
5. Close with a clear call to action.
6. Keep to 3-4 paragraphs, max 350 words.
7. Never use clichés like "I am a hard worker", "team player", or "passionate about".
8. Do NOT include salutation address lines or sign-off — return body paragraphs only.
"""


# ── Prompt builder ────────────────────────────────────────────────────────────

def _build_resume_snippet(parsed: dict) -> str:
    lines: list[str] = []
    skills = [s.get("name", s) if isinstance(s, dict) else s for s in parsed.get("skills", [])]
    if skills:
        lines.append(f"Skills: {', '.join(skills[:20])}")
    for e in parsed.get("experience", [])[:3]:
        title    = e.get("title", "")
        company  = e.get("company", "")
        bullets  = e.get("bullets") or e.get("description", "")
        duration = e.get("duration", "")
        bullet_text = ""
        if isinstance(bullets, list):
            bullet_text = " | ".join(bullets[:3])
        elif isinstance(bullets, str):
            bullet_text = bullets[:200]
        lines.append(f"Role: {title} at {company} ({duration}) — {bullet_text}")
    for ed in parsed.get("education", [])[:1]:
        lines.append(f"Education: {ed.get('degree','')} from {ed.get('institution','')}")
    return "\n".join(lines)


def build_prompt(
    parsed_resume: dict,
    job_description: str,
    company_name: str,
    job_title: str,
    tone: Tone,
) -> str:
    resume_snippet = _build_resume_snippet(parsed_resume)
    tone_note = {
        "formal":   "formal and professional — polished, third-person gravitas",
        "friendly": "warm and conversational — like a message to someone you admire",
        "direct":   "punchy and confident — brevity over ceremony, results over pleasantries",
    }[tone]

    return textwrap.dedent(f"""
        ---
        TASK: Write a cover letter body for the following:

        Target Role: {sanitize_input(job_title, 200)} at {sanitize_input(company_name, 200)}
        Tone: {tone_note}

        Candidate Resume Summary:
        {resume_snippet}

        Job Description (excerpt):
        {sanitize_input(job_description, 1500)}

        Write the cover letter body paragraphs now (no salutation, no sign-off):
    """).strip()


# ── Generation ────────────────────────────────────────────────────────────────

def generate_cover_letter(
    parsed_resume: dict,
    job_description: str,
    company_name: str,
    job_title: str,
    tone: Tone = "formal",
) -> str:
    system_inst = f"{_SYSTEM}\n\n{_FEW_SHOTS}"
    model  = genai.GenerativeModel(_MODEL_NAME, system_instruction=system_inst)
    prompt = build_prompt(parsed_resume, job_description, company_name, job_title, tone)
    config = genai.types.GenerationConfig(temperature=0.7, max_output_tokens=600)
    resp   = model.generate_content(prompt, generation_config=config)
    return resp.text.strip()


# ── Export helpers ────────────────────────────────────────────────────────────

def export_pdf(
    letter_body: str,
    candidate_name: str,
    job_title: str,
    company_name: str,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=15,
        spaceAfter=8,
    )

    story = [
        Paragraph(candidate_name, heading_style),
        Spacer(1, 0.3 * cm),
        Paragraph(f"{job_title} — {company_name}", body_style),
        Spacer(1, 0.5 * cm),
    ]

    for para in letter_body.split("\n\n"):
        para = para.strip()
        if para:
            story.append(Paragraph(para.replace("\n", " "), body_style))

    doc.build(story)
    return buf.getvalue()


def export_docx(
    letter_body: str,
    candidate_name: str,
    job_title: str,
    company_name: str,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Header
    h = doc.add_heading(candidate_name, level=1)
    h.runs[0].font.size = Pt(13)

    sub = doc.add_paragraph(f"{job_title} — {company_name}")
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].italic    = True
    doc.add_paragraph()

    # Body
    for para in letter_body.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para.replace("\n", " "))
            p.paragraph_format.space_after = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
